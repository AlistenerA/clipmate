"""
MCP File Reader Server
桥接智谱 GLM-4V API，提供图片/PDF/Word 文件解析能力
解析结果以纯文本形式返回，供 DeepSeek 等模型使用
"""
import asyncio
import io
import logging
import tempfile
from pathlib import Path

from mcp.server import Server
from mcp.server.stdio import stdio_server
import mcp.types as types

from common import (
    DEFAULT_VISION_PROMPT,
    DEFAULT_PDF_NATIVE_PROMPT,
    SUPPORTED_IMG,
    call_glm_native_file,
    call_glm_vision,
    format_glm_response,
    get_client,
    parse_page_range,
    pdf_page_to_vision_text,
    validate_file_path,
    validate_image_path,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("file-reader")

server = Server("file-reader")

get_client()


def _to_bool(val) -> bool:
    """将字符串 'true'/'false' 或布尔值统一转换为 bool"""
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.lower() in ("true", "1", "yes")
    return bool(val)


@server.list_tools()
async def list_tools() -> list[types.Tool]:
    return [
        types.Tool(
            name="read_image",
            description=(
                "【必须使用】读取和分析图片文件（.png / .jpg / .jpeg / .webp / .gif / .bmp）。"
                "当需要查看、识别、描述、分析图片内容时，必须调用本工具。"
                "因为 deepseek-v4-pro 等模型无法直接处理图片输入，"
                "本工具会通过视觉模型将图片转为文字描述后返回。"
                "参数 file_path 为图片文件的绝对路径。"
                "返回图片的文字描述、OCR 文字、图表数据等。"
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "图片文件的绝对路径"
                    },
                    "prompt": {
                        "type": "string",
                        "description": "可选的提示词",
                        "default": DEFAULT_VISION_PROMPT,
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="read_pdf",
            description=(
                "读取 PDF 文件。三种模式："
                "1) 默认：本地 pdfplumber 提取文本（快速、免费）；"
                "2) use_vision=True：将每页转为图片，用 GLM-4V 视觉理解（适合扫描件/图表）；"
                "3) use_native=True：将整个 PDF 以 file_url 直接发送给 GLM-4.6V 原生理解（适合复杂文档/跨页逻辑）。"
                "use_native 无视 page_range，总是处理整份文档。"
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "PDF 文件的绝对路径"
                    },
                    "page_range": {
                        "type": "string",
                        "description": "页面范围，如 '1-3' 或 '1,3,5'，默认 'all'",
                        "default": "all"
                    },
                    "use_vision": {
                        "type": "boolean",
                        "description": "使用 GLM-4V 逐页视觉理解（适合扫描件/图表）",
                        "default": False
                    },
                    "use_native": {
                        "type": "boolean",
                        "description": "使用 GLM-4.6V 原生文档理解（整个 PDF 直接送入模型）",
                        "default": False
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="read_docx",
            description="读取 Word (.docx) 文件，提取全部文本内容。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "DOCX 文件的绝对路径"
                    }
                },
                "required": ["file_path"]
            }
        )
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[types.TextContent]:
    try:
        if name == "read_image":
            return await _handle_read_image(arguments)
        elif name == "read_pdf":
            return await _handle_read_pdf(arguments)
        elif name == "read_docx":
            return await _handle_read_docx(arguments)
        else:
            return [types.TextContent(type="text", text=f"Unknown tool: {name}")]
    except FileNotFoundError as e:
        return [types.TextContent(type="text", text=str(e))]
    except ValueError as e:
        return [types.TextContent(type="text", text=str(e))]
    except Exception as e:
        logger.exception("Unexpected error in tool %s", name)
        return [types.TextContent(type="text", text=f"Error: {e}")]


async def _handle_read_image(args: dict) -> list[types.TextContent]:
    file_path = args["file_path"]
    prompt = args.get("prompt", DEFAULT_VISION_PROMPT)

    validate_image_path(file_path)

    result = await asyncio.to_thread(call_glm_vision, [file_path], prompt)
    text = format_glm_response(
        content=result["content"],
        reasoning_content=result.get("reasoning_content", ""),
        usage=result.get("usage"),
    )
    return [types.TextContent(type="text", text=text)]


async def _handle_read_pdf(args: dict) -> list[types.TextContent]:
    file_path = args["file_path"]
    page_range = args.get("page_range", "all")
    use_vision = _to_bool(args.get("use_vision", False))
    use_native = _to_bool(args.get("use_native", False))

    p = Path(file_path)
    if not p.exists():
        return [types.TextContent(type="text", text=f"File not found: {file_path}")]

    if use_native:
        return await _handle_read_pdf_native(file_path, page_range)

    import pdfplumber

    with pdfplumber.open(str(p)) as pdf:
        total = len(pdf.pages)
        if page_range == "all":
            indices = list(range(total))
        else:
            try:
                indices = parse_page_range(page_range, total)
            except ValueError as e:
                return [types.TextContent(type="text", text=str(e))]

        if not indices:
            return [types.TextContent(type="text", text="No pages matched the given range.")]

        results = []
        if use_vision:
            tmpdir = tempfile.gettempdir()
            for i in indices:
                text = await asyncio.to_thread(
                    pdf_page_to_vision_text, pdf, i, total, tmpdir
                )
                results.append(text)
        else:
            for i in indices:
                text = pdf.pages[i].extract_text()
                if text and text.strip():
                    results.append(f"=== Page {i + 1}/{total} ===\n{text}")
                else:
                    results.append(
                        f"=== Page {i + 1}/{total} ===\n"
                        f"(no extractable text, try use_vision=true or use_native=true)"
                    )

        output = "\n\n".join(results)
        return [types.TextContent(type="text", text=output)]


async def _handle_read_pdf_native(
    file_path: str, page_range: str = "all"
) -> list[types.TextContent]:
    validate_file_path(file_path)

    prompt = DEFAULT_PDF_NATIVE_PROMPT
    if page_range != "all":
        prompt = (
            f"{DEFAULT_PDF_NATIVE_PROMPT} "
            f"请重点查看第 {page_range} 页的内容。"
        )

    result = await asyncio.to_thread(call_glm_native_file, file_path, prompt)
    text = format_glm_response(
        content=result["content"],
        reasoning_content=result.get("reasoning_content", ""),
        usage=result.get("usage"),
        source_label="[GLM-4.6V Native Document Understanding]",
    )
    return [types.TextContent(type="text", text=text)]


async def _handle_read_docx(args: dict) -> list[types.TextContent]:
    file_path = args["file_path"]

    p = Path(file_path)
    if not p.exists():
        return [types.TextContent(type="text", text=f"File not found: {file_path}")]

    import docx

    doc = docx.Document(str(p))
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    for table in doc.tables:
        table_text = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            table_text.append(" | ".join(cells))
        paragraphs.append("\n" + "\n".join(table_text) + "\n")

    text = "\n\n".join(paragraphs)
    return [types.TextContent(type="text", text=text)]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )


if __name__ == "__main__":
    import sys
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(
            asyncio.WindowsProactorEventLoopPolicy()
        )
    asyncio.run(main())
