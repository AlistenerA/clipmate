"""
read_file.py — 独立 CLI 工具，桥接智谱 GLM-4V / GLM-4 API
当 MCP server 不可用时，Agent 可通过 bash 调用本脚本来读取文件

用法:
  python read_file.py image   <文件路径>
  python read_file.py pdf     <文件路径> [--vision] [--native]
  python read_file.py docx    <文件路径>
"""
import argparse
import sys
import tempfile
from pathlib import Path

from common import (
    DEFAULT_VISION_PROMPT,
    DEFAULT_PDF_NATIVE_PROMPT,
    call_glm_native_file,
    call_glm_vision,
    format_glm_response,
    parse_page_range,
    pdf_page_to_vision_text,
    validate_file_path,
    validate_image_path,
)


def read_image(file_path: str, prompt: str | None = None) -> str:
    validate_image_path(file_path)
    result = call_glm_vision([file_path], prompt or DEFAULT_VISION_PROMPT)
    return format_glm_response(
        content=result["content"],
        reasoning_content=result.get("reasoning_content", ""),
        usage=result.get("usage"),
    )


def read_pdf(file_path: str, use_vision: bool = False, use_native: bool = False, page_range: str = "all") -> str:
    p = Path(file_path)
    if not p.exists():
        return f"ERROR: File not found: {file_path}"

    if use_native:
        return read_pdf_native(file_path, page_range)

    try:
        import pdfplumber
    except ImportError:
        return "ERROR: pdfplumber not installed"

    with pdfplumber.open(str(p)) as pdf:
        total = len(pdf.pages)
        if page_range == "all":
            indices = list(range(total))
        else:
            try:
                indices = parse_page_range(page_range, total)
            except ValueError as e:
                return f"ERROR: {e}"

        if not indices:
            return "ERROR: No pages matched the given range."

        results = []
        if use_vision:
            tmpdir = tempfile.gettempdir()
            for i in indices:
                text = pdf_page_to_vision_text(pdf, i, total, tmpdir)
                results.append(text)
        else:
            for i in indices:
                text = pdf.pages[i].extract_text()
                if text and text.strip():
                    results.append(f"=== Page {i + 1}/{total} ===\n{text}")
                else:
                    results.append(
                        f"=== Page {i + 1}/{total} ===\n(no extractable text, try --vision or --native)"
                    )

    return "\n\n".join(results)


def read_pdf_native(file_path: str, page_range: str = "all") -> str:
    validate_file_path(file_path)

    prompt = DEFAULT_PDF_NATIVE_PROMPT
    if page_range != "all":
        prompt = (
            f"{DEFAULT_PDF_NATIVE_PROMPT} "
            f"请重点查看第 {page_range} 页的内容。"
        )

    result = call_glm_native_file(file_path, prompt)
    return format_glm_response(
        content=result["content"],
        reasoning_content=result.get("reasoning_content", ""),
        usage=result.get("usage"),
        source_label="[GLM-4.6V Native Document Understanding]",
    )


def read_docx(file_path: str) -> str:
    p = Path(file_path)
    if not p.exists():
        return f"ERROR: File not found: {file_path}"

    try:
        import docx
    except ImportError:
        return "ERROR: python-docx not installed"

    doc = docx.Document(str(p))
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    for table in doc.tables:
        table_text = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            table_text.append(" | ".join(cells))
        paragraphs.append("\n" + "\n".join(table_text) + "\n")

    return "\n\n".join(paragraphs)


def main():
    parser = argparse.ArgumentParser(description="Read files via ZhipuAI GLM")
    sub = parser.add_subparsers(dest="command", required=True)

    img = sub.add_parser("image", help="Read an image file")
    img.add_argument("path", help="Image file path")

    pdf = sub.add_parser("pdf", help="Read a PDF file")
    pdf.add_argument("path", help="PDF file path")
    pdf.add_argument("--vision", action="store_true", help="Use GLM-4V vision for scanned PDFs")
    pdf.add_argument("--native", action="store_true", help="Use GLM-4.6V native document understanding")

    docx = sub.add_parser("docx", help="Read a DOCX file")
    docx.add_argument("path", help="DOCX file path")

    args = parser.parse_args()

    try:
        if args.command == "image":
            result = read_image(args.path)
        elif args.command == "pdf":
            result = read_pdf(
                args.path,
                use_vision=getattr(args, "vision", False),
                use_native=getattr(args, "native", False),
            )
        elif args.command == "docx":
            result = read_docx(args.path)
        else:
            result = f"Unknown command: {args.command}"
    except FileNotFoundError as e:
        result = f"ERROR: {e}"
    except ValueError as e:
        result = f"ERROR: {e}"
    except Exception as e:
        result = f"ERROR: {e}"

    sys.stdout.reconfigure(encoding="utf-8")
    sys.stdout.buffer.write((result + "\n").encode("utf-8", errors="replace"))


if __name__ == "__main__":
    main()
